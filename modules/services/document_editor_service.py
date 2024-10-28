# services/document_editor_service.py

import logging
import threading
from typing import Any, Dict, List, Optional
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from modules.utilities.logging_manager import setup_logging
from modules.utilities.config_loader import ConfigLoader
from modules.security.encryption_manager import EncryptionManager
from modules.security.authentication import AuthenticationManager


class DocumentEditorServiceError(Exception):
    """Custom exception for DocumentEditorService-related errors."""
    pass


class DocumentEditorService:
    """
    Provides document editing capabilities, including creating, reading, updating, and formatting
    Microsoft Word documents. Utilizes the python-docx library to ensure comprehensive document handling.
    Ensures secure handling of files and configurations.
    """

    def __init__(self):
        """
        Initializes the DocumentEditorService with necessary configurations and authentication.
        """
        self.logger = setup_logging('DocumentEditorService')
        self.config_loader = ConfigLoader()
        self.encryption_manager = EncryptionManager()
        self.auth_manager = AuthenticationManager()
        self.lock = threading.Lock()
        self.logger.info("DocumentEditorService initialized successfully.")

    def create_document(self, output_path: str, title: str, author: str, content: List[str], images: Optional[List[str]] = None) -> bool:
        """
        Creates a new Microsoft Word document with specified content and images.

        Args:
            output_path (str): The file path for the created document.
            title (str): The title of the document.
            author (str): The author of the document.
            content (List[str]): A list of paragraphs to include in the document.
            images (Optional[List[str]], optional): A list of image file paths to embed. Defaults to None.

        Returns:
            bool: True if the document is created successfully, False otherwise.
        """
        try:
            self.logger.debug(f"Creating document at '{output_path}' with title '{title}', author '{author}', and {len(content)} paragraphs.")
            document = Document()
            document.core_properties.title = title
            document.core_properties.author = author

            # Add title
            title_paragraph = document.add_heading(title, level=1)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            self.logger.debug("Title added to the document.")

            # Add content
            for para in content:
                paragraph = document.add_paragraph(para)
                self.logger.debug(f"Added paragraph: {para}")

            # Add images
            if images:
                for image_path in images:
                    if os.path.exists(image_path):
                        document.add_picture(image_path, width=Inches(4))
                        last_paragraph = document.paragraphs[-1]
                        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        self.logger.debug(f"Added image '{image_path}' to the document.")
                    else:
                        self.logger.warning(f"Image file '{image_path}' does not exist. Skipping.")

            # Save document
            document.save(output_path)
            self.logger.info(f"Document created successfully at '{output_path}'.")
            return True
        except Exception as e:
            self.logger.error(f"Error creating document at '{output_path}': {e}", exc_info=True)
            return False

    def read_document(self, doc_path: str) -> Optional[Dict[str, Any]]:
        """
        Reads a Microsoft Word document and extracts its content and metadata.

        Args:
            doc_path (str): The file path to the document.

        Returns:
            Optional[Dict[str, Any]]: A dictionary containing document metadata and content, or None if reading fails.
        """
        try:
            self.logger.debug(f"Reading document from '{doc_path}'.")
            if not os.path.exists(doc_path):
                self.logger.error(f"Document '{doc_path}' does not exist.")
                return None
            document = Document(doc_path)
            metadata = {
                'title': document.core_properties.title,
                'author': document.core_properties.author,
                'created': document.core_properties.created.isoformat() if document.core_properties.created else None,
                'last_modified': document.core_properties.modified.isoformat() if document.core_properties.modified else None
            }
            content = [para.text for para in document.paragraphs]
            self.logger.info(f"Document '{doc_path}' read successfully.")
            return {'metadata': metadata, 'content': content}
        except Exception as e:
            self.logger.error(f"Error reading document '{doc_path}': {e}", exc_info=True)
            return None

    def update_document(self, doc_path: str, new_content: Optional[List[str]] = None, new_images: Optional[List[str]] = None) -> bool:
        """
        Updates an existing Microsoft Word document by adding new content and images.

        Args:
            doc_path (str): The file path to the document to update.
            new_content (Optional[List[str]], optional): A list of new paragraphs to add. Defaults to None.
            new_images (Optional[List[str]], optional): A list of new image file paths to embed. Defaults to None.

        Returns:
            bool: True if the document is updated successfully, False otherwise.
        """
        try:
            self.logger.debug(f"Updating document '{doc_path}' with {len(new_content) if new_content else 0} new paragraphs and {len(new_images) if new_images else 0} new images.")
            if not os.path.exists(doc_path):
                self.logger.error(f"Document '{doc_path}' does not exist.")
                return False
            document = Document(doc_path)

            # Add new content
            if new_content:
                for para in new_content:
                    paragraph = document.add_paragraph(para)
                    self.logger.debug(f"Added new paragraph: {para}")

            # Add new images
            if new_images:
                for image_path in new_images:
                    if os.path.exists(image_path):
                        document.add_picture(image_path, width=Inches(4))
                        last_paragraph = document.paragraphs[-1]
                        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        self.logger.debug(f"Added new image '{image_path}' to the document.")
                    else:
                        self.logger.warning(f"Image file '{image_path}' does not exist. Skipping.")

            # Save changes
            document.save(doc_path)
            self.logger.info(f"Document '{doc_path}' updated successfully.")
            return True
        except Exception as e:
            self.logger.error(f"Error updating document '{doc_path}': {e}", exc_info=True)
            return False

    def format_document(self, doc_path: str, style_changes: Dict[str, Any]) -> bool:
        """
        Applies formatting changes to a Microsoft Word document.

        Args:
            doc_path (str): The file path to the document.
            style_changes (Dict[str, Any]): A dictionary containing style changes (e.g., font size, alignment).

        Returns:
            bool: True if the document is formatted successfully, False otherwise.
        """
        try:
            self.logger.debug(f"Formatting document '{doc_path}' with style changes: {style_changes}.")
            if not os.path.exists(doc_path):
                self.logger.error(f"Document '{doc_path}' does not exist.")
                return False
            document = Document(doc_path)

            for paragraph in document.paragraphs:
                if 'font_size' in style_changes:
                    for run in paragraph.runs:
                        run.font.size = style_changes['font_size']
                if 'alignment' in style_changes:
                    if style_changes['alignment'].lower() == 'center':
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    elif style_changes['alignment'].lower() == 'left':
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    elif style_changes['alignment'].lower() == 'right':
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    elif style_changes['alignment'].lower() == 'justify':
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Apply additional styles if provided
            if 'styles' in style_changes:
                for style_name, style_props in style_changes['styles'].items():
                    if style_props['type'] == 'paragraph':
                        style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                        style.font.name = style_props.get('font_name', 'Arial')
                        style.font.size = style_props.get('font_size', 12)
                        style.paragraph_format.alignment = getattr(WD_PARAGRAPH_ALIGNMENT, style_props.get('alignment', 'LEFT').upper())
                    elif style_props['type'] == 'character':
                        style = document.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
                        style.font.name = style_props.get('font_name', 'Arial')
                        style.font.size = style_props.get('font_size', 12)
                        style.font.bold = style_props.get('bold', False)
                        style.font.italic = style_props.get('italic', False)

            # Save formatted document
            document.save(doc_path)
            self.logger.info(f"Document '{doc_path}' formatted successfully.")
            return True
        except Exception as e:
            self.logger.error(f"Error formatting document '{doc_path}': {e}", exc_info=True)
            return False

    def add_table_to_document(self, doc_path: str, table_data: List[List[Any]], table_style: Optional[str] = 'Light List Accent 1') -> bool:
        """
        Adds a table to a Microsoft Word document.

        Args:
            doc_path (str): The file path to the document.
            table_data (List[List[Any]]): A list of rows, each containing a list of cell data.
            table_style (Optional[str], optional): The style to apply to the table. Defaults to 'Light List Accent 1'.

        Returns:
            bool: True if the table is added successfully, False otherwise.
        """
        try:
            self.logger.debug(f"Adding table to document '{doc_path}' with style '{table_style}'.")
            if not os.path.exists(doc_path):
                self.logger.error(f"Document '{doc_path}' does not exist.")
                return False
            document = Document(doc_path)

            if not table_data:
                self.logger.warning("No table data provided. Skipping table addition.")
                return False

            table = document.add_table(rows=1, cols=len(table_data[0]))
            table.style = table_style

            # Add header row
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(table_data[0]):
                hdr_cells[i].text = str(header)
            self.logger.debug("Header row added to the table.")

            # Add data rows
            for row_data in table_data[1:]:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = str(cell_data)
            self.logger.debug(f"Added {len(table_data) - 1} rows to the table.")

            # Save document
            document.save(doc_path)
            self.logger.info(f"Table added successfully to '{doc_path}'.")
            return True
        except Exception as e:
            self.logger.error(f"Error adding table to document '{doc_path}': {e}", exc_info=True)
            return False

    def close_service(self):
        """
        Closes any resources or sessions held by the service.
        """
        try:
            self.logger.debug("Closing DocumentEditorService resources.")
            # Currently, no persistent resources to close
            self.logger.info("DocumentEditorService closed successfully.")
        except Exception as e:
            self.logger.error(f"Error closing DocumentEditorService: {e}", exc_info=True)
            raise DocumentEditorServiceError(f"Error closing DocumentEditorService: {e}")
